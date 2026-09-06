import csv
from pathlib import Path

import pandas as pd
from pypdf import PdfReader
from docx import Document

from text_cleaner import clean_text


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".pdf",
    ".docx",
    ".csv"
}


def read_txt(file_path):
    encodings = [
        "utf-8",
        "utf-8-sig",
        "latin-1"
    ]

    for encoding in encodings:
        try:
            return Path(
                file_path
            ).read_text(
                encoding=encoding
            )
        except UnicodeDecodeError:
            continue

    raise ValueError(
        "Unable to decode the TXT file."
    )


def read_pdf(file_path):
    reader = PdfReader(
        file_path
    )

    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1
    ):
        text = page.extract_text()

        if text:
            pages.append(
                f"[Page {page_number}]\n{text}"
            )

    return "\n\n".join(pages)


def read_docx(file_path):
    document = Document(
        file_path
    )

    sections = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    for table_number, table in enumerate(
        document.tables,
        start=1
    ):
        sections.append(
            f"[Table {table_number}]"
        )

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            sections.append(
                " | ".join(cells)
            )

    return "\n".join(sections)


def read_csv(file_path):
    dataframe = pd.read_csv(
        file_path
    )

    dataframe = dataframe.fillna("")

    lines = []

    for _, row in dataframe.iterrows():
        values = []

        for column in dataframe.columns:
            value = str(
                row[column]
            ).strip()

            if value:
                values.append(
                    f"{column}: {value}"
                )

        if values:
            lines.append(
                " | ".join(values)
            )

    return "\n".join(lines)


def extract_text(file_path):
    path = Path(
        file_path
    )

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. "
            "Use TXT, PDF, DOCX or CSV."
        )

    if extension == ".txt":
        text = read_txt(path)

    elif extension == ".pdf":
        text = read_pdf(path)

    elif extension == ".docx":
        text = read_docx(path)

    elif extension == ".csv":
        text = read_csv(path)

    else:
        raise ValueError(
            "Unsupported file type."
        )

    return clean_text(text)